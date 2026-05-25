from __future__ import annotations
from pathlib import Path
from typing import Protocol, runtime_checkable

from loguru import logger

SUPPORTED_EXTENSIONS = {".txt", ".md", ".rst", ".pdf", ".docx", ".html", ".htm"}


@runtime_checkable
class Reader(Protocol):
    def can_read(self, path: Path) -> bool: ...
    def read(self, path: Path) -> str: ...


# ---------------------------------------------------------------------------
# Plain text (TXT / MD / RST)
# ---------------------------------------------------------------------------

class TextReader:
    EXTENSIONS = {".txt", ".md", ".rst", ".log"}

    def can_read(self, path: Path) -> bool:
        return path.suffix.lower() in self.EXTENSIONS

    def read(self, path: Path) -> str:
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            try:
                import chardet
                raw = path.read_bytes()
                enc = chardet.detect(raw).get("encoding") or "latin-1"
                return raw.decode(enc, errors="replace")
            except Exception:
                return path.read_text(encoding="latin-1", errors="replace")


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

class PDFReader:
    EXTENSIONS = {".pdf"}

    def can_read(self, path: Path) -> bool:
        return path.suffix.lower() in self.EXTENSIONS

    def read(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            parts: list[str] = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(f"[Page {i + 1}]\n{text}")
            return "\n\n".join(parts)
        except Exception as exc:
            logger.warning(f"PDF read failed for {path.name}: {exc}")
            return ""


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

class DocxReader:
    EXTENSIONS = {".docx"}

    def can_read(self, path: Path) -> bool:
        return path.suffix.lower() in self.EXTENSIONS

    def read(self, path: Path) -> str:
        try:
            from docx import Document
            doc = Document(str(path))
            parts: list[str] = []

            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    # Preserve heading level as markdown
                    style = para.style.name if para.style else ""
                    if "Heading" in style:
                        level = "".join(c for c in style if c.isdigit()) or "1"
                        parts.append(f"{'#' * int(level)} {t}")
                    else:
                        parts.append(t)

            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    parts.append("\n".join(rows))

            return "\n\n".join(parts)
        except Exception as exc:
            logger.warning(f"DOCX read failed for {path.name}: {exc}")
            return ""


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

class HtmlReader:
    EXTENSIONS = {".html", ".htm"}

    def can_read(self, path: Path) -> bool:
        return path.suffix.lower() in self.EXTENSIONS

    def read(self, path: Path) -> str:
        try:
            from bs4 import BeautifulSoup
            raw = path.read_bytes()
            soup = BeautifulSoup(raw, "lxml")
            for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
                tag.decompose()
            return soup.get_text(separator="\n", strip=True)
        except Exception as exc:
            logger.warning(f"HTML read failed for {path.name}: {exc}")
            return ""


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

_READERS: list[Reader] = [
    TextReader(),
    PDFReader(),
    DocxReader(),
    HtmlReader(),
]


def read_document(path: Path) -> str:
    """Read a document and return its plain-text content."""
    for reader in _READERS:
        if reader.can_read(path):
            return reader.read(path)
    raise ValueError(f"Unsupported file type: {path.suffix!r}")


def is_supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_EXTENSIONS
